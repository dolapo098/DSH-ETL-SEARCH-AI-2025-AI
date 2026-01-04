import io

import docx

from app.contracts.providers.i_document_text_extractor import IDocumentTextExtractor


class WordDocumentExtractor(IDocumentTextExtractor):
    def extract_text(self, file_content: bytes) -> str:
        doc = docx.Document(io.BytesIO(file_content))
        parts = []

        for p in doc.paragraphs:
            
            if p.text.strip():
                parts.append(p.text)

        for table in doc.tables:
            
            for row in table.rows:
                
                for cell in row.cells:
                    
                    if cell.text.strip():
                        parts.append(cell.text)

        return "\n".join(parts).strip()

